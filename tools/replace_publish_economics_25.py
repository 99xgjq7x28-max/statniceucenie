import json
import re
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SOURCE = Path("/Users/michaljeck/Downloads/VYPRACOVANE_EKONOMIA A FINANCIE.docx")
PUBLISH = ROOT / "statniceucenie-publish"
TARGET = PUBLISH / "data/topics.json"


def clean(text):
    return re.sub(r"\s+", " ", text).strip()


def build_offline(source_name, script_name, output_name, topics):
    html = (PUBLISH / source_name).read_text(encoding="utf-8")
    css_match = re.search(
        r'<link rel="stylesheet" href="src/styles\.css\?v=(\d+)">',
        html,
    )
    script_match = re.search(
        rf'<script src="src/{re.escape(script_name)}\?v=(\d+)" defer></script>',
        html,
    )
    if css_match is None or script_match is None:
        raise RuntimeError(f"V {source_name} sa nenašli verziované CSS/JS odkazy.")

    css = (PUBLISH / "src/styles.css").read_text(encoding="utf-8")
    script = (PUBLISH / "src" / script_name).read_text(encoding="utf-8")
    embedded = json.dumps(topics, ensure_ascii=False, indent=2)
    html = html.replace(css_match.group(0), f"<style>\n{css}\n</style>")
    html = html.replace(
        script_match.group(0),
        f"<script>window.__TOPICS__ = {embedded};\n{script}\n</script>",
    )
    (PUBLISH / output_name).write_text(html, encoding="utf-8")


document = Document(SOURCE)
start = next(
    index
    for index, paragraph in enumerate(document.paragraphs)
    if clean(paragraph.text) == "25."
)
end = next(
    index
    for index in range(start + 1, len(document.paragraphs))
    if clean(document.paragraphs[index].text) == "26."
)
paragraphs = [
    paragraph
    for paragraph in document.paragraphs[start:end]
    if clean(paragraph.text)
]

title = clean(paragraphs[1].text)
blocks = []

for paragraph in paragraphs[2:]:
    text = clean(paragraph.text)
    properties = paragraph._p.pPr
    is_list = properties is not None and properties.numPr is not None

    if re.match(r"^[a-d]\)\s", text) or text == "Manažérsky prínos: Zachytenie vzťahov":
        blocks.append({"type": "heading", "text": text})
    elif text.startswith("Záver:"):
        blocks.append({"type": "heading", "text": "Záver"})
        blocks.append({"type": "p", "text": clean(text.removeprefix("Záver:"))})
    elif is_list:
        blocks.append({"type": "li", "text": text})
    else:
        blocks.append({"type": "p", "text": text})

with TARGET.open(encoding="utf-8") as file:
    topics = json.load(file)

topic = next((item for item in topics if item.get("id") == "ekonomia-25"), None)
if topic is None:
    raise RuntimeError("V topics.json sa nenašla otázka ekonomia-25.")

topic["title"] = title
topic["blocks"] = blocks
topic["wordCount"] = sum(len(block["text"].split()) for block in blocks)
topic["tableCount"] = 0
topic["imageCount"] = 0
topic["cue"] = (
    "Podvojné účtovníctvo zachytáva každú transakciu najmenej na dvoch účtoch. "
    "Otázka spája odberateľsko-dodávateľské vzťahy, cenné papiere, mzdy a dane "
    "s riadením likvidity, nákladov a vzťahov podniku s okolím."
)
topic["emergencyStart"] = [
    (
        "Podvojné účtovníctvo je založené na princípe podvojnosti, takže každá "
        "hospodárska operácia ovplyvňuje najmenej dva účty. Finančné účty a "
        "zúčtovacie vzťahy zachytávajú interné aj externé finančné toky podniku."
    ),
    (
        "Medzi hlavné oblasti patria pohľadávky voči odberateľom, záväzky voči "
        "dodávateľom, majetkové a dlhové cenné papiere, mzdová agenda a priame "
        "či nepriame dane."
    ),
    (
        "Pre manažment sú tieto informácie dôležité pri riadení likvidity, "
        "kontrole nákladov, hodnotení platobnej disciplíny partnerov a "
        "dodržiavaní povinností voči zamestnancom, poisťovniam a štátu."
    ),
]

with TARGET.open("w", encoding="utf-8") as file:
    json.dump(topics, file, ensure_ascii=False, indent=2)
    file.write("\n")

build_offline("index.html", "app.js", "offline.html", topics)
build_offline("plan.html", "plan.js", "plan-offline.html", topics)

print(
    json.dumps(
        {
            "id": topic["id"],
            "blocks": len(topic["blocks"]),
            "words": topic["wordCount"],
            "tables": topic["tableCount"],
            "images": topic["imageCount"],
        },
        ensure_ascii=False,
    )
)
