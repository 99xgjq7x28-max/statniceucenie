import hashlib
import json
import re
import shutil
import unicodedata
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "VYPRACOVANE_MANAZMENT_OPRAVENE.docx"
OFFICIAL = Path(
    "/Users/michaljeck/Library/Mobile Documents/com~apple~CloudDocs/"
    "Univerzita Komenského/3.ročník_3Mbx23_X3_2025:2026/Statnice/"
    "Otazky_SS Bc_Mbx_Manazment 25_26.docx"
)
DATA = ROOT / "data/topics.json"
ASSETS = ROOT / "assets/images"


def normalize(text):
    value = unicodedata.normalize("NFKC", text).replace("\xa0", " ")
    return re.sub(r"\s+", " ", value).strip()


def comparison_text(text):
    value = unicodedata.normalize("NFKD", normalize(text))
    value = value.encode("ascii", "ignore").decode().lower()
    return re.sub(r"[^a-z0-9]+", " ", value).strip()


def iter_body(document):
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def official_questions():
    document = Document(OFFICIAL)
    questions = [
        normalize(paragraph.text)
        for paragraph in document.paragraphs
        if paragraph.style.name == "Body Text Indent 2" and normalize(paragraph.text)
    ]
    if len(questions) != 30:
        raise RuntimeError(f"Expected 30 official questions, found {len(questions)}.")
    return questions


def paragraph_images(paragraph, document, topic_number, image_index):
    blocks = []
    for blip in paragraph._p.xpath(".//a:blip"):
        relationship_id = blip.get(qn("r:embed"))
        if not relationship_id:
            continue
        part = document.part.related_parts[relationship_id]
        extension = Path(part.partname).suffix.lower() or ".png"
        digest = hashlib.sha1(part.blob).hexdigest()[:10]
        filename = f"manazment-v2-q{topic_number:02d}-{image_index:02d}-{digest}{extension}"
        destination = ASSETS / filename
        destination.write_bytes(part.blob)
        blocks.append(
            {
                "type": "image",
                "src": f"assets/images/{filename}",
                "alt": f"Manažment, otázka {topic_number} – obrázok {image_index}",
            }
        )
        image_index += 1
    return blocks, image_index


def is_heading(paragraph, text):
    style = paragraph.style.name or ""
    if style.startswith("Heading"):
        return True
    runs = [run for run in paragraph.runs if normalize(run.text)]
    if runs and all(run.bold is True for run in runs) and len(text) <= 180:
        return True
    if re.match(r"^\d+\.\s+\S", text) and len(text) <= 160:
        return True
    return False


def is_list_item(paragraph, text):
    properties = paragraph._p.pPr
    if properties is not None and properties.numPr is not None:
        return True
    if "List" in (paragraph.style.name or ""):
        return True
    return bool(re.match(r"^[•▪●◦\-–]\s+", text))


def extract_topics():
    document = Document(SOURCE)
    questions = official_questions()
    question_lookup = {comparison_text(question): index + 1 for index, question in enumerate(questions)}
    topics = {
        number: {
            "subject": "Manažment",
            "number": number,
            "title": questions[number - 1],
            "blocks": [],
            "id": f"manazment-{number}",
        }
        for number in range(1, 31)
    }

    current = None
    image_indices = {number: 1 for number in range(1, 31)}
    matched = []

    for element in iter_body(document):
        if isinstance(element, Paragraph):
            text = normalize(element.text)
            match = question_lookup.get(comparison_text(text)) if text else None
            if match:
                current = match
                matched.append(match)
                continue
            if current is None:
                continue
            if re.fullmatch(r"\d+\.", text):
                continue

            image_blocks, image_indices[current] = paragraph_images(
                element, document, current, image_indices[current]
            )
            if text:
                block_type = "p"
                if is_heading(element, text):
                    block_type = "heading"
                elif is_list_item(element, text):
                    block_type = "li"
                topics[current]["blocks"].append({"type": block_type, "text": text})
            topics[current]["blocks"].extend(image_blocks)
        elif current is not None:
            rows = [
                [normalize(cell.text) for cell in row.cells]
                for row in element.rows
            ]
            if any(any(cell for cell in row) for row in rows):
                topics[current]["blocks"].append({"type": "table", "rows": rows})

    if matched != list(range(1, 31)):
        raise RuntimeError(f"Questions are missing or out of order: {matched}")

    for topic in topics.values():
        text_blocks = [
            block["text"]
            for block in topic["blocks"]
            if block["type"] in {"p", "heading", "li"}
        ]
        table_text = [
            cell
            for block in topic["blocks"]
            if block["type"] == "table"
            for row in block["rows"]
            for cell in row
        ]
        topic["wordCount"] = sum(len(text.split()) for text in text_blocks + table_text)
        topic["tableCount"] = sum(block["type"] == "table" for block in topic["blocks"])
        topic["imageCount"] = sum(block["type"] == "image" for block in topic["blocks"])

        prose = [
            block["text"]
            for block in topic["blocks"]
            if block["type"] in {"p", "li"} and len(block["text"].split()) >= 8
        ]
        headings = [
            block["text"]
            for block in topic["blocks"]
            if block["type"] == "heading"
        ]
        topic["cue"] = prose[0][:500] if prose else topic["title"]
        emergency = prose[:2]
        if headings:
            emergency.append("Hlavné časti odpovede: " + "; ".join(headings[:5]) + ".")
        topic["emergencyStart"] = emergency[:3]

    return [topics[number] for number in range(1, 31)]


def verify_import(topics):
    if len(topics) != 30:
        raise RuntimeError("The import does not contain 30 management topics.")
    if [topic["number"] for topic in topics] != list(range(1, 31)):
        raise RuntimeError("Management topic numbering is incomplete.")
    if sum(topic["tableCount"] for topic in topics) != 4:
        raise RuntimeError("Not all 4 tables were imported.")
    if sum(topic["imageCount"] for topic in topics) != 4:
        raise RuntimeError("Not all 4 images were imported.")
    if any(not topic["blocks"] for topic in topics):
        raise RuntimeError("At least one topic has an empty answer.")
    product_owner = " ".join(
        block.get("text", "")
        for block in topics[16]["blocks"]
    )
    required = [
        "Product Owner pritom nie je automaticky projektovým manažérom",
        "Vývojári participatívne rozhodujú",
    ]
    if not all(text in product_owner for text in required):
        raise RuntimeError("The Product Owner correction is missing from question 17.")


def update_data(management_topics):
    existing = json.loads(DATA.read_text(encoding="utf-8"))
    economics = [topic for topic in existing if topic["subject"] == "Ekonómia a financie"]
    if len(economics) != 30:
        raise RuntimeError(f"Expected 30 economics topics, found {len(economics)}.")
    combined = economics + management_topics
    DATA.write_text(
        json.dumps(combined, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    return combined


def build_offline(source_html, script_path, output_path, topics):
    html = (ROOT / source_html).read_text(encoding="utf-8")
    css_link = '<link rel="stylesheet" href="src/styles.css?v=19">'
    script_tag = f'<script src="{script_path}?v=19" defer></script>'
    if css_link not in html or script_tag not in html:
        raise RuntimeError(f"Versioned assets not found in {source_html}.")
    css = (ROOT / "src/styles.css").read_text(encoding="utf-8")
    script = (ROOT / script_path).read_text(encoding="utf-8")
    embedded = json.dumps(topics, ensure_ascii=False, indent=2)
    html = html.replace(css_link, f"<style>\n{css}\n</style>")
    html = html.replace(
        script_tag,
        f"<script>window.__TOPICS__ = {embedded};\n{script}\n</script>",
    )
    output_path.write_text(html, encoding="utf-8")


def update_publish(topics):
    for directory in [ROOT / "publish-statniceucenie", ROOT / "statniceucenie-publish"]:
        (directory / "data").mkdir(parents=True, exist_ok=True)
        (directory / "src").mkdir(parents=True, exist_ok=True)
        (directory / "assets/images").mkdir(parents=True, exist_ok=True)
        for filename in ["index.html", "plan.html", "offline.html", "plan-offline.html", "README.md"]:
            shutil.copy2(ROOT / filename, directory / filename)
        for filename in ["app.js", "plan.js", "styles.css"]:
            shutil.copy2(ROOT / "src" / filename, directory / "src" / filename)
        shutil.copy2(DATA, directory / "data/topics.json")
        for image in ASSETS.iterdir():
            if image.is_file():
                shutil.copy2(image, directory / "assets/images" / image.name)


management = extract_topics()
verify_import(management)
all_topics = update_data(management)
build_offline("index.html", "src/app.js", ROOT / "offline.html", all_topics)
build_offline("plan.html", "src/plan.js", ROOT / "plan-offline.html", all_topics)
update_publish(all_topics)

print(
    json.dumps(
        {
            "topics": len(management),
            "words": sum(topic["wordCount"] for topic in management),
            "tables": sum(topic["tableCount"] for topic in management),
            "images": sum(topic["imageCount"] for topic in management),
        },
        ensure_ascii=False,
    )
)
