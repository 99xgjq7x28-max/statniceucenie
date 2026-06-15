import collections
import difflib
import json
import re
import unicodedata
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SOURCE = Path(
    "/Users/michaljeck/Library/Mobile Documents/com~apple~CloudDocs/"
    "Univerzita Komenského/3.ročník_3Mbx23_X3_2025:2026/Statnice/"
    "vypracovane /Ekonómia a financie statnice.docx"
)
OFFICIAL = Path(
    "/Users/michaljeck/Library/Mobile Documents/com~apple~CloudDocs/"
    "Univerzita Komenského/3.ročník_3Mbx23_X3_2025:2026/Statnice/"
    "Otazky_SS Bc_Mbx_Ekonomia a financie 25_26.docx"
)
DATA = ROOT / "data/topics.json"


def clean(text):
    value = unicodedata.normalize("NFKC", text).replace("\xa0", " ")
    return re.sub(r"\s+", " ", value).strip()


def normalized_words(text):
    value = unicodedata.normalize("NFKD", clean(text))
    value = value.encode("ascii", "ignore").decode().lower()
    return re.sub(r"[^a-z0-9]+", " ", value).strip().split()


source_paragraphs = [
    clean(paragraph.text)
    for paragraph in Document(SOURCE).paragraphs
    if clean(paragraph.text)
]
official_questions = [
    clean(paragraph.text)
    for paragraph in Document(OFFICIAL).paragraphs[11:41]
]
topics = [
    topic
    for topic in json.loads(DATA.read_text(encoding="utf-8"))
    if topic["subject"] == "Ekonómia a financie"
]

if len(official_questions) != 30 or len(topics) != 30:
    raise RuntimeError("The audit requires exactly 30 official and 30 app questions.")

starts = []
header_scores = []
for question in official_questions:
    question_words = normalized_words(question)
    scores = [
        difflib.SequenceMatcher(
            None, question_words, normalized_words(paragraph)
        ).ratio()
        if len(paragraph) > 60
        else 0
        for paragraph in source_paragraphs
    ]
    index = max(range(len(scores)), key=scores.__getitem__)
    starts.append(index)
    header_scores.append(scores[index])

if starts != sorted(starts) or len(set(starts)) != 30:
    raise RuntimeError(f"Question headers are missing or out of order: {starts}")

results = []
for offset, (start, topic) in enumerate(zip(starts, topics)):
    end = starts[offset + 1] if offset + 1 < 30 else len(source_paragraphs)
    source_words = []
    for paragraph in source_paragraphs[start + 1:end]:
        source_words.extend(normalized_words(paragraph))

    app_words = []
    for block in topic["blocks"]:
        if "text" in block:
            app_words.extend(normalized_words(block["text"]))
        elif block["type"] == "table":
            for row in block["rows"]:
                for cell in row:
                    app_words.extend(normalized_words(cell))

    source_counts = collections.Counter(source_words)
    app_counts = collections.Counter(app_words)
    missing = sum((source_counts - app_counts).values())
    extra = sum((app_counts - source_counts).values())
    sequence_ratio = difflib.SequenceMatcher(
        None, source_words, app_words
    ).ratio()
    results.append(
        {
            "number": offset + 1,
            "sourceWords": len(source_words),
            "appWords": len(app_words),
            "sequenceRatio": round(sequence_ratio, 6),
            "missingWords": missing,
            "extraWords": extra,
            "headerScore": round(header_scores[offset], 6),
        }
    )

failures = [
    result
    for result in results
    if result["sequenceRatio"] != 1
    or result["missingWords"]
    or result["extraWords"]
]
if failures:
    raise RuntimeError(
        "Economics content mismatch: "
        + json.dumps(failures, ensure_ascii=False)
    )

media = {
    "tables": sum(topic["tableCount"] for topic in topics),
    "images": sum(topic["imageCount"] for topic in topics),
}
if media != {"tables": 0, "images": 0}:
    raise RuntimeError(f"Unexpected economics media counts: {media}")

print(
    json.dumps(
        {
            "questions": len(results),
            "exactMatches": len(results) - len(failures),
            "sourceWords": sum(item["sourceWords"] for item in results),
            "appWords": sum(item["appWords"] for item in results),
            "media": media,
            "results": results,
        },
        ensure_ascii=False,
    )
)
