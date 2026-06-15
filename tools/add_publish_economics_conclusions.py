import json
import re
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SOURCE = Path("/Users/michaljeck/Downloads/VYPRACOVANE_EKONOMIA A FINANCIE.docx")
PUBLISH = ROOT / "statniceucenie-publish"
TARGET = PUBLISH / "data/topics.json"


def clean(text):
    return re.sub(r"\s+", " ", text).strip()


def comparable(text):
    return clean(text).replace("–", "-").replace("—", "-")


def question_ranges(document):
    markers = []
    for index, paragraph in enumerate(document.paragraphs):
        text = clean(paragraph.text)
        match = re.fullmatch(r"(\d{1,2})\.", text)
        if match and 1 <= int(match.group(1)) <= 30:
            markers.append((int(match.group(1)), index))

    if [number for number, _ in markers] != list(range(1, 31)):
        raise RuntimeError("V dokumente sa nepodarilo jednoznačne nájsť otázky 1 až 30.")

    ranges = {}
    for position, (number, start) in enumerate(markers):
        end = markers[position + 1][1] if position + 1 < len(markers) else len(document.paragraphs)
        ranges[number] = (start, end)
    return ranges


def extract_end_conclusions(document):
    conclusions = {}
    ranges = question_ranges(document)

    for number, (start, end) in ranges.items():
        paragraphs = [
            (index, clean(document.paragraphs[index].text))
            for index in range(start, end)
            if clean(document.paragraphs[index].text)
        ]

        for position, (_, text) in enumerate(paragraphs):
            if re.fullmatch(r"Záver", text, re.IGNORECASE):
                if position + 1 >= len(paragraphs):
                    raise RuntimeError(f"Otázka E{number} má prázdny záver.")
                conclusions[number] = paragraphs[position + 1][1]
            elif re.match(r"^Záver(?: pre manažéra)?:\s*.+", text, re.IGNORECASE):
                conclusions[number] = re.sub(
                    r"^Záver(?: pre manažéra)?:\s*",
                    "",
                    text,
                    flags=re.IGNORECASE,
                )

    return conclusions


def build_offline(source_name, script_name, output_name, topics):
    html = (PUBLISH / source_name).read_text(encoding="utf-8")
    css_match = re.search(
        r'<link rel="stylesheet" href="src/styles\.css\?v=(\d+)">',
        html,
    )
    script_match = re.search(
        rf'<script src="src/{re.escape(script_name)}\?v=(\d+)" defer></script>',
        html,
    )
    if css_match is None or script_match is None:
        raise RuntimeError(f"V {source_name} sa nenašli verziované CSS/JS odkazy.")

    css = (PUBLISH / "src/styles.css").read_text(encoding="utf-8")
    script = (PUBLISH / "src" / script_name).read_text(encoding="utf-8")
    embedded = json.dumps(topics, ensure_ascii=False, indent=2)
    html = html.replace(css_match.group(0), f"<style>\n{css}\n</style>")
    html = html.replace(
        script_match.group(0),
        f"<script>window.__TOPICS__ = {embedded};\n{script}\n</script>",
    )
    (PUBLISH / output_name).write_text(html, encoding="utf-8")


document = Document(SOURCE)
conclusions = extract_end_conclusions(document)

with TARGET.open(encoding="utf-8") as file:
    topics = json.load(file)

changed = []
skipped = []

for number, conclusion in sorted(conclusions.items()):
    topic = next(item for item in topics if item.get("id") == f"ekonomia-{number}")
    existing_texts = [
        clean(block.get("text", ""))
        for block in topic["blocks"]
        if block.get("type") in {"heading", "p", "li"}
    ]

    existing_comparable = {comparable(text) for text in existing_texts}
    if (
        comparable(conclusion) in existing_comparable
        or comparable(f"Záver: {conclusion}") in existing_comparable
    ):
        skipped.append(number)
        continue

    topic["blocks"].extend(
        [
            {"type": "heading", "text": "Záver"},
            {"type": "p", "text": conclusion},
        ]
    )
    topic["wordCount"] += len("Záver".split()) + len(conclusion.split())
    changed.append(number)

with TARGET.open("w", encoding="utf-8") as file:
    json.dump(topics, file, ensure_ascii=False, indent=2)
    file.write("\n")

build_offline("index.html", "app.js", "offline.html", topics)
build_offline("plan.html", "plan.js", "plan-offline.html", topics)

print(
    json.dumps(
        {
            "source_conclusions": sorted(conclusions),
            "added": changed,
            "already_present": skipped,
        },
        ensure_ascii=False,
    )
)
